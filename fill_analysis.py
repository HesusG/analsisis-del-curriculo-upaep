#!/usr/bin/env python3
"""
fill_analysis.py — Llena el template "Análisis de Planeación.docx"
con el análisis del curso MT1001B del Tecnológico de Monterrey.

Fuentes: 6 archivos .md en sources/
Salida:  "Análisis de Planeación - MT1001B.docx"
"""

from docx import Document
from docx.shared import Pt, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import copy


# ── Helpers ────────────────────────────────────────────────────────

FONT_NAME = "Arial"
FONT_SIZE = Pt(12)


def _set_run_format(run, bold=False, italic=False):
    """Apply consistent formatting to a run."""
    run.font.name = FONT_NAME
    run.font.size = FONT_SIZE
    run.bold = bold
    run.italic = italic


def write_paragraph(para, text, bold=False, italic=False):
    """Clear a paragraph and write new text with formatting."""
    for run in para.runs:
        run.text = ""
    run = para.add_run(text)
    _set_run_format(run, bold=bold, italic=italic)
    return run


def append_to_paragraph(para, text, bold=False, italic=False):
    """Append text to an existing paragraph (after existing content)."""
    run = para.add_run(text)
    _set_run_format(run, bold=bold, italic=italic)
    return run


def unmerge_observation_column(table):
    """Remove vertical merge from observation column (col 3) in a table.

    Table 0 has rows 1-5 merged into one cell for observations.
    This unmerges them so each row can have its own observation.
    """
    for row in table.rows:
        tr = row._tr
        tcs = tr.findall(qn('w:tc'))
        if len(tcs) >= 4:
            tc = tcs[3]
            tcPr = tc.find(qn('w:tcPr'))
            if tcPr is not None:
                vm_el = tcPr.find(qn('w:vMerge'))
                if vm_el is not None:
                    tcPr.remove(vm_el)


def _fill_cell_by_xml(table, row_idx, col_idx, text):
    """Fill a cell using raw XML tc access (avoids merged-cell confusion)."""
    tr = table.rows[row_idx]._tr
    tcs = tr.findall(qn('w:tc'))
    tc = tcs[col_idx]
    # Get or create the first paragraph
    p_elements = tc.findall(qn('w:p'))
    if p_elements:
        p_el = p_elements[0]
    else:
        return
    # Clear existing runs
    for r in p_el.findall(qn('w:r')):
        p_el.remove(r)
    # Create new run with text
    from docx.text.paragraph import Paragraph
    para = Paragraph(p_el, tc)
    run = para.add_run(text)
    _set_run_format(run)


def fill_table_row(table, row_idx, si=True, observacion=""):
    """Fill a checklist table row: mark Sí or No and add observation."""
    si_text = "Sí" if si else ""
    no_text = "No" if not si else ""

    _fill_cell_by_xml(table, row_idx, 1, si_text)
    _fill_cell_by_xml(table, row_idx, 2, no_text)
    _fill_cell_by_xml(table, row_idx, 3, observacion)


def insert_paragraphs_after(doc, anchor_para, texts):
    """Insert new paragraphs after a given paragraph element.

    Each item in `texts` can be:
      - str: plain paragraph
      - tuple(str, dict): text + kwargs for _set_run_format (bold, italic)
    """
    # Get the XML element of the anchor
    anchor_elem = anchor_para._element
    parent = anchor_elem.getparent()

    inserted = []
    for item in texts:
        if isinstance(item, tuple):
            text, kwargs = item
        else:
            text, kwargs = item, {}

        # Create a new paragraph by copying format from anchor
        new_p = copy.deepcopy(anchor_para._element)
        # Clear all runs
        for r in new_p.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}r'):
            new_p.remove(r)

        # Insert after current last inserted (or anchor)
        ref = inserted[-1] if inserted else anchor_elem
        ref.addnext(new_p)
        inserted.append(new_p)

        # Now create a proper paragraph wrapper and add the run
        from docx.text.paragraph import Paragraph
        para_obj = Paragraph(new_p, anchor_para._element.getparent())
        run = para_obj.add_run(text)
        _set_run_format(run, **kwargs)

    return inserted


# ── Main fill functions ────────────────────────────────────────────

def fill_metadata(doc):
    """P4-P6: Evaluador, Institución, Responsable."""
    paras = doc.paragraphs

    # P4: Nombre del Evaluador
    append_to_paragraph(paras[4], " HesusG")

    # P5: Institución que analiza
    append_to_paragraph(paras[5], " UPAEP (Universidad Popular Autónoma del Estado de Puebla)")

    # P6: Responsable de la Planeación
    append_to_paragraph(
        paras[6],
        " No especificado — documento institucional del bloque MT1001B, "
        "Tecnológico de Monterrey"
    )


def fill_parte1_contexto(doc):
    """Parte 1 (P13-P30): Diagnóstico del contexto institucional y estudiantil."""
    paras = doc.paragraphs

    contenido = [
        # Contexto institucional
        (
            "Contexto Institucional",
            {"bold": True}
        ),
        (
            "El bloque MT1001B «Descubrimientos del mercado para el desarrollo de "
            "estrategias» se inscribe en el Modelo Tec21 del Tecnológico de Monterrey, "
            "un modelo educativo basado en aprendizaje experiencial y retos. Este modelo "
            "privilegia la formación por competencias, la vinculación con socios formadores "
            "del sector productivo y la flexibilidad curricular, lo que posiciona a la "
            "institución en un paradigma de educación orientada a la empleabilidad y la "
            "innovación."
        ),
        (
            "Desde la perspectiva de Gimeno Sacristán (2010), el currículum prescrito del "
            "Tec21 expresa un ideal formativo particular: un profesionista con competencias "
            "técnicas de negocio y capacidad de innovación. Sin embargo, al analizar el "
            "currículum operacional —es decir, lo que efectivamente sucede en el aula—, "
            "emergen brechas entre la intención y la práctica que este análisis buscará "
            "identificar."
        ),
        (
            "La planeación revisada corresponde al Módulo 1: «Diagnóstico y desarrollo "
            "de la estrategia», que abarca 20 horas de clase y 10 horas de tarea, "
            "distribuidas en 10 sesiones a lo largo de aproximadamente 4 semanas. El "
            "bloque completo (60 horas) se organiza en tres módulos de 20 horas cada uno."
        ),
        "",  # spacer
        # Perfil estudiantil
        (
            "Perfil del Estudiante",
            {"bold": True}
        ),
        (
            "Los participantes son estudiantes de segundo semestre de la Entrada de "
            "Negocios, en una etapa temprana de su trayectoria de exploración profesional. "
            "Su nivel de dominio esperado oscila entre A y B:"
        ),
        (
            "• Nivel A: El estudiante realiza tareas en situaciones controladas o simuladas, "
            "requiere acompañamiento docente y tiene poca autonomía."
        ),
        (
            "• Nivel B: El estudiante atiende tareas con mayor autonomía y complejidad "
            "creciente, asumiendo algunos riesgos y tomando decisiones."
        ),
        (
            "Este perfil implica que la planeación debe ofrecer andamiaje suficiente para "
            "guiar al estudiante de nivel A, al tiempo que proporcione espacios de creciente "
            "autonomía para facilitar la transición hacia el nivel B. La secuencia didáctica, "
            "como se verá más adelante, cumple parcialmente con este requisito."
        ),
        "",  # spacer
        (
            "Necesidades Identificadas",
            {"bold": True}
        ),
        (
            "El bloque busca desarrollar tres competencias clave: SNEG0400 (Inteligencia "
            "de negocios), SNEG0700 (Mercados y oportunidades de negocio) y SEG0200 "
            "(Emprendimiento innovador). Estas competencias responden a la necesidad del "
            "sector empresarial de profesionistas capaces de diagnosticar mercados, formular "
            "estrategias y proponer soluciones innovadoras con base en datos."
        ),
        (
            "No obstante, desde el enfoque de Apple (1979), es pertinente preguntarse: "
            "¿para quién y para qué se diseña este currículum? La orientación del bloque "
            "está alineada con las demandas del mercado y los socios formadores, lo cual "
            "es comprensible en un contexto de negocios, pero deja poco espacio para la "
            "reflexión crítica sobre las implicaciones éticas de las estrategias de "
            "mercadotecnia o el impacto social del consumo."
        ),
    ]

    # Fill P13 onwards
    start_idx = 13
    for i, item in enumerate(contenido):
        idx = start_idx + i
        if idx >= len(paras):
            break
        if isinstance(item, tuple):
            text, kwargs = item
            write_paragraph(paras[idx], text, **kwargs)
        elif item == "":
            # Leave empty as spacer
            write_paragraph(paras[idx], "")
        else:
            write_paragraph(paras[idx], item)


def fill_parte2_datos(doc):
    """Parte 2: Table 0 — Datos de presentación."""
    t = doc.tables[0]

    # Table 0 has observation column vertically merged — unmerge it first
    unmerge_observation_column(t)

    fill_table_row(t, 1, si=True,
        observacion=(
            "Se identifica implícitamente como Tecnológico de Monterrey a "
            "través del Modelo Tec21 y la estructura de bloques. No aparece "
            "el nombre completo de la institución de forma explícita en el "
            "encabezado de la planeación."
        ))

    fill_table_row(t, 2, si=False,
        observacion=(
            "No se especifica el periodo o semestre lectivo concreto. "
            "Se recomienda incluir el año y periodo académico para "
            "facilitar el seguimiento curricular."
        ))

    fill_table_row(t, 3, si=False,
        observacion=(
            "No se menciona el nombre del docente titular responsable "
            "de la asignatura. Dado que es un documento institucional "
            "de diseño del bloque, es comprensible, pero se recomienda "
            "incluir al responsable de implementación."
        ))

    fill_table_row(t, 4, si=True,
        observacion=(
            "Se identifica claramente: MT1001B «Descubrimientos del "
            "mercado para el desarrollo de estrategias», 2° semestre, "
            "Entrada de Negocios. Incluye clave del bloque y ubicación "
            "en la malla curricular."
        ))

    fill_table_row(t, 5, si=True,
        observacion=(
            "Se incluye duración total (60 horas, 5 semanas), "
            "distribución en 3 módulos de 20 horas cada uno, y "
            "desglose de horas clase (20h) vs. tarea (10h) para "
            "el Módulo 1."
        ))


def fill_parte3_proposito(doc):
    """Parte 3: Table 1 — Propósito u objetivo general + recomendaciones P51."""
    t = doc.tables[1]

    fill_table_row(t, 1, si=True,
        observacion=(
            "El propósito se articula con estructura Qué/Cómo/Para qué: "
            "descubrir el mercado (qué), mediante un reto con plan de "
            "mercadotecnia (cómo), para hacer recomendaciones innovadoras "
            "al socio formador (para qué). La redacción es funcional, "
            "aunque podría beneficiarse de mayor concisión."
        ))

    fill_table_row(t, 2, si=True,
        observacion=(
            "El objetivo es claro: formular un plan básico de "
            "mercadotecnia a partir de un diagnóstico de mercado. "
            "El alcance está bien delimitado al nivel de las "
            "competencias A y B."
        ))

    fill_table_row(t, 3, si=True,
        observacion=(
            "El propósito es alcanzable en las 5 semanas (60 horas) "
            "previstas. La secuenciación en 3 módulos (diagnóstico, "
            "análisis, reto) permite un avance progresivo. Las 10 "
            "sesiones del Módulo 1 distribuyen adecuadamente la carga."
        ))

    fill_table_row(t, 4, si=True,
        observacion=(
            "El resultado es observable y evaluable: se entrega un plan "
            "de mercadotecnia y formatos de design thinking. Los 5 "
            "objetos de aprendizaje tienen puntajes asignados "
            "(5+5+2+2+6=20 puntos)."
        ))

    # P51: Recomendaciones
    paras = doc.paragraphs
    write_paragraph(paras[51],
        "El propósito cumple satisfactoriamente con los criterios evaluados. "
        "Como recomendación menor, se sugiere explicitar en la redacción del "
        "objetivo el vínculo con las competencias transversales (pensamiento "
        "crítico, colaboración) que el propio modelo Tec21 promueve. Desde la "
        "perspectiva de multiliteracidades (Cope & Kalantzis, 2009), también "
        "sería deseable que el propósito mencionara explícitamente el uso de "
        "múltiples modos de representación (visual, digital, textual) como "
        "parte del proceso formativo, no solo como productos de entrega."
    )


def fill_parte4_competencias(doc):
    """Parte 4: Tables 2-3 — Competencias y aprendizajes esperados + narrativa."""
    paras = doc.paragraphs

    # Table 2: Competencias
    t2 = doc.tables[2]

    fill_table_row(t2, 1, si=True,
        observacion=(
            "Las competencias SNEG0400, SNEG0700 y SEG0200 presentan "
            "estructura de verbo + contenido + finalidad. Ejemplo: "
            "«Analizar información cuantitativa y cualitativa (verbo + "
            "contenido) para la toma de decisiones (finalidad)». La "
            "condición de referencia es parcialmente explícita en los "
            "niveles de dominio A y B."
        ))

    fill_table_row(t2, 2, si=True,
        observacion=(
            "Las competencias son claras y concisas. Cada una incluye "
            "subcompetencias específicas con descriptores diferenciados "
            "por nivel de dominio."
        ))

    fill_table_row(t2, 3, si=True,
        observacion=(
            "Son observables y evaluables a través de los objetos de "
            "aprendizaje: diagnóstico de caso (SNEG0700), evaluación "
            "conceptual (SNEG0400), plan de mercadotecnia (SNEG0703) "
            "y taller de design thinking (SEG0201)."
        ))

    # Table 3: Aprendizajes esperados
    t3 = doc.tables[3]

    fill_table_row(t3, 1, si=True,
        observacion=(
            "Los 5 objetos de aprendizaje expresan con claridad lo "
            "esperado: diagnóstico de mercadotecnia, dominio conceptual, "
            "identificación de ejes rectores, elaboración de plan y "
            "aplicación de design thinking."
        ))

    fill_table_row(t3, 2, si=False,
        observacion=(
            "Los objetos de aprendizaje están diseñados para niveles A-B, "
            "donde el estudiante requiere guía. La autonomía se favorece "
            "parcialmente: el OA4 (guía para plan de MKT) pide investigación "
            "autónoma, pero otros objetos son predominantemente guiados. Se "
            "recomienda incorporar actividades de mayor agencia estudiantil."
        ))

    fill_table_row(t3, 3, si=True,
        observacion=(
            "Se expresan conocimientos (conceptos de MKT, segmentación), "
            "habilidades (diagnóstico, visualización, design thinking) y "
            "actitudes (proactividad ante soluciones). Los contenidos "
            "actitudinales, sin embargo, son más declarativos que "
            "operacionales."
        ))

    fill_table_row(t3, 4, si=True,
        observacion=(
            "Los verbos son precisos y alineados con los niveles de "
            "dominio: identificar (A), analizar (A-B), diseñar (A), "
            "generar (A). Reflejan una taxonomía coherente para segundo "
            "semestre."
        ))

    fill_table_row(t3, 5, si=True,
        observacion=(
            "Los aprendizajes son relevantes para la competencia de "
            "mercados, claros en su formulación y evaluables mediante "
            "entregas específicas con puntajes asignados."
        ))

    # Narrativa P61-P62
    write_paragraph(paras[61],
        "Las competencias del bloque MT1001B están bien estructuradas y "
        "alineadas con el perfil de egreso de la Entrada de Negocios. La "
        "diferenciación por niveles de dominio (A y B) es un acierto que "
        "permite calibrar las expectativas según la etapa formativa del "
        "estudiante. Sin embargo, las competencias se enfocan casi exclusivamente "
        "en habilidades técnicas de negocio, sin integrar explícitamente "
        "competencias de pensamiento crítico o ciudadanía digital que, desde "
        "el conectivismo (Siemens, 2005), serían fundamentales para un "
        "profesionista que opera en entornos de información distribuida."
    )


def fill_parte5_contenidos(doc):
    """Parte 5: Table 4 + narrativa — Contenidos y subtemas."""
    paras = doc.paragraphs
    t4 = doc.tables[4]

    fill_table_row(t4, 1, si=True,
        observacion=(
            "Los contenidos están directamente relacionados con la "
            "asignatura de mercadotecnia: orientación al consumidor, "
            "propuesta de valor, segmentación, mezcla de mercadotecnia, "
            "plan de MKT, design thinking."
        ))

    fill_table_row(t4, 2, si=True,
        observacion=(
            "Los contenidos aportan al logro del propósito. La secuencia "
            "diagnóstico → conceptos → estrategia → innovación → evaluación "
            "construye progresivamente hacia el plan de mercadotecnia."
        ))

    fill_table_row(t4, 3, si=True,
        observacion=(
            "Los contenidos favorecen la praxis: cada sesión teórica "
            "(clase del profesor) se acompaña de una actividad aplicada "
            "(análisis de caso, plenarias, talleres). El balance "
            "conceptual-procedimental es adecuado."
        ))

    fill_table_row(t4, 4, si=True,
        observacion=(
            "Los contenidos son adecuados para estudiantes de 2° semestre "
            "con niveles A-B. Parten de conceptos básicos (orientación al "
            "consumidor) y avanzan hacia aplicaciones más complejas (plan "
            "de MKT, design thinking)."
        ))

    # Narrativa P69-P70
    write_paragraph(paras[69],
        "El análisis de los contenidos del Módulo 1 revela un equilibrio "
        "adecuado entre temas conceptuales y procedimentales. Las 10 "
        "sesiones cubren desde fundamentos de mercadotecnia hasta la "
        "aplicación creativa con design thinking, siguiendo un orden lógico "
        "que permite al estudiante construir conocimiento de forma progresiva. "
        "Los contenidos procedimentales (elaboración de reportes, identificación "
        "de oportunidades) y actitudinales (proactividad) están presentes, "
        "aunque estos últimos se mencionan de forma genérica."
    )

    write_paragraph(paras[70],
        "Desde el enfoque de Apple (1979), los contenidos seleccionados "
        "privilegian una visión funcionalista de la mercadotecnia centrada "
        "en satisfacer necesidades del mercado y generar valor para la "
        "organización. No se observan contenidos que inviten a cuestionar "
        "críticamente las prácticas de consumo, la ética publicitaria o "
        "el impacto socioambiental de las estrategias de mercado. Incorporar "
        "al menos un módulo de reflexión crítica enriquecería la formación "
        "integral del estudiante, alineándose con lo que Sacristán (2010) "
        "denomina el análisis del «contenido cultural» que el currículum "
        "decide transmitir."
    )


def fill_parte6_secuencia(doc):
    """Parte 6 (P75-P80): Secuencia didáctica."""
    paras = doc.paragraphs

    write_paragraph(paras[75],
        "La secuencia didáctica del Módulo 1 se organiza mediante un modelo "
        "pedagógico de 6 fases cíclicas —Empatiza, Explica, Explora, Elabora, "
        "Innova y Evalúa— que constituye una adaptación del modelo 5E "
        "(Bybee et al.) enriquecida con una fase de innovación. Esta "
        "estructura facilita un aprendizaje experiencial coherente con el "
        "Modelo Tec21."
    )

    write_paragraph(paras[76],
        "En cuanto a las fases didácticas tradicionales: la fase de Inicio "
        "se cubre con Empatiza (presentación del caso, identificación de la "
        "situación); el Procesamiento se desarrolla a través de Explica y "
        "Explora (contenidos conceptuales y análisis de información); el "
        "Reforzamiento corresponde a Elabora e Innova (reportes de hallazgos "
        "y propuestas innovadoras); y el Cierre se concreta en Evalúa "
        "(evaluación de aprendizajes con retroalimentación continua). Las "
        "cuatro fases quedan cubiertas satisfactoriamente."
    )

    write_paragraph(paras[77],
        "Un aspecto positivo es que la retroalimentación se declara como "
        "«continua a lo largo de todo el proceso», lo que sugiere un enfoque "
        "formativo. Sin embargo, la planeación no especifica los mecanismos "
        "concretos de esta retroalimentación: ¿se utilizan rúbricas? "
        "¿retroalimentación entre pares? ¿herramientas digitales? Esta "
        "ambigüedad debilita la operacionalización de un principio que, en "
        "sí mismo, es pedagógicamente valioso."
    )

    write_paragraph(paras[78],
        "Desde la perspectiva de Sacristán (2010), la secuencia muestra una "
        "brecha entre el currículum prescrito y el operacional: el modelo "
        "cíclico se describe a nivel conceptual pero no se detalla cómo "
        "se implementa en cada sesión. Las actividades en la tabla del "
        "cronograma (clases magistrales, plenarias, talleres) no siempre "
        "se corresponden explícitamente con las fases del modelo 5E "
        "adaptado. Se recomienda que la planeación vincule cada sesión "
        "con la fase del modelo que le corresponde, para facilitar tanto "
        "la implementación como la observación de la práctica docente."
    )

    write_paragraph(paras[79],
        "Como recomendación adicional, la secuencia podría beneficiarse "
        "de incorporar momentos de metacognición explícita donde los "
        "estudiantes reflexionen sobre su propio proceso de aprendizaje, "
        "lo que fortalecería la autonomía progresiva que el tránsito de "
        "nivel A a B requiere."
    )


def fill_parte7_metodologia(doc):
    """Parte 7: Table 5 — Metodología."""
    t5 = doc.tables[5]

    fill_table_row(t5, 1, si=True,
        observacion=(
            "Las estrategias de enseñanza del docente son claras: clases "
            "magistrales sobre conceptos clave (orientación al consumidor, "
            "propuesta de valor, segmentación, plan de MKT), presentación "
            "de casos y conducción de plenarias y talleres."
        ))

    fill_table_row(t5, 2, si=True,
        observacion=(
            "Las estrategias de aprendizaje del estudiante están definidas: "
            "análisis de mini caso, elaboración de diagnósticos, investigación "
            "secundaria, participación en plenarias, llenado de formatos de "
            "design thinking y elaboración de reportes."
        ))

    fill_table_row(t5, 3, si=True,
        observacion=(
            "Las estrategias son claras y secuenciadas. Cada sesión combina "
            "un momento expositivo del docente con una actividad aplicada "
            "del estudiante."
        ))

    fill_table_row(t5, 4, si=True,
        observacion=(
            "Las estrategias son adecuadas para estudiantes de nivel A-B: "
            "parten de la exposición guiada y avanzan hacia la aplicación "
            "con acompañamiento (plenarias) y la creación (design thinking)."
        ))

    fill_table_row(t5, 5, si=False,
        observacion=(
            "Las estrategias son funcionales pero no innovadoras. "
            "Predominan métodos tradicionales (clase magistral + caso "
            "de estudio). El taller de design thinking es el elemento "
            "más innovador, pero se limita a la última sesión. Desde "
            "el enfoque de multiliteracidades (Cope & Kalantzis, 2009), "
            "falta integración de recursos multimodales, herramientas "
            "digitales y actividades que involucren múltiples modos de "
            "representación del conocimiento."
        ))

    fill_table_row(t5, 6, si=True,
        observacion=(
            "Las estrategias son coherentes con el perfil de segundo "
            "semestre. Sin embargo, podrían diversificarse incluyendo "
            "aprendizaje colaborativo digital, simulaciones de mercado "
            "y portafolios electrónicos, alineándose con el conectivismo "
            "(Siemens, 2005)."
        ))

    # Narrativa P85-P86
    paras = doc.paragraphs
    write_paragraph(paras[85],
        "La metodología del Módulo 1 combina enseñanza expositiva con "
        "aprendizaje basado en casos y design thinking. Esta combinación "
        "es coherente con el enfoque del Modelo Tec21. Sin embargo, la "
        "simulación Wideband Delphi realizada como parte de este análisis "
        "otorgó una calificación de 5.5/10 en integración tecnológica, "
        "señalando que las estrategias dependen excesivamente de métodos "
        "presenciales tradicionales sin aprovechar herramientas digitales "
        "que potenciarían la experiencia de aprendizaje."
    )


def fill_parte8_evaluacion(doc):
    """Parte 8: Table 6 + narrativa — Evaluación."""
    t6 = doc.tables[6]

    fill_table_row(t6, 1, si=True,
        observacion=(
            "Los criterios de evaluación están vinculados a las estrategias: "
            "el diagnóstico de caso (OA1, 5 pts), la evaluación conceptual "
            "(OA2, 5 pts), los ejes rectores (OA3, 2 pts), la guía del "
            "plan de MKT (OA4, 2 pts) y el taller de design thinking "
            "(OA5, 6 pts) corresponden a las actividades desarrolladas "
            "en la secuencia didáctica."
        ))

    fill_table_row(t6, 2, si=False,
        observacion=(
            "La puntuación no está equilibrada de forma equitativa. El OA5 "
            "(design thinking) vale 6 puntos mientras que OA3 y OA4 valen "
            "solo 2 cada uno. Si bien es válido ponderar más las actividades "
            "integradoras, la brecha es significativa. Además, falta "
            "evaluación diagnóstica explícita y la evaluación actitudinal "
            "no tiene instrumentos específicos asignados."
        ))

    # Narrativa P90-P91
    paras = doc.paragraphs
    write_paragraph(paras[90],
        "El sistema de evaluación del Módulo 1 contempla 5 objetos de "
        "aprendizaje con un total de 20 puntos. Se observa evaluación "
        "formativa a través de las actividades en clase y la retroalimentación "
        "continua declarada en el modelo 5E adaptado, y evaluación sumativa "
        "mediante entregas con puntajes asignados. Sin embargo, no se "
        "identifica una evaluación diagnóstica explícita al inicio del módulo "
        "que permita identificar conocimientos previos de los estudiantes."
    )

    write_paragraph(paras[91],
        "La evaluación se concentra en productos conceptuales y procedimentales "
        "(reportes, exámenes, planes), pero no incluye instrumentos específicos "
        "para evaluar los contenidos actitudinales declarados (proactividad "
        "ante soluciones). Se recomienda incorporar autoevaluación, evaluación "
        "entre pares y rúbricas que incluyan criterios actitudinales, "
        "fortaleciendo la coherencia entre lo que se enseña, lo que se "
        "practica y lo que se evalúa."
    )


def fill_parte9_recursos(doc):
    """Parte 9 (P95-P98): Análisis de recursos."""
    paras = doc.paragraphs

    write_paragraph(paras[95],
        "Los recursos mencionados en la planeación incluyen: casos de "
        "estudio empresariales, bibliografía de referencia (Lamb, Hair y "
        "McDaniel para conceptos de mercadotecnia), formatos impresos "
        "para el taller de design thinking y productos físicos con "
        "empaque y etiqueta para el ejercicio de la sesión 8. Estos "
        "recursos son adecuados para los objetivos del módulo y "
        "accesibles para estudiantes de segundo semestre."
    )

    write_paragraph(paras[96],
        "Sin embargo, se identifican ausencias significativas en materia "
        "de recursos tecnológicos y digitales. La planeación no menciona: "
        "plataformas de análisis de datos (Tableau, Power BI, Google "
        "Analytics), herramientas de colaboración digital (Miro, Notion, "
        "Google Workspace), simuladores de mercado, bases de datos de "
        "investigación de mercados, ni recursos multimedia o audiovisuales. "
        "Esta carencia fue identificada consistentemente en la simulación "
        "Wideband Delphi (calificación de 5.5/10 en integración tecnológica), "
        "donde los cinco expertos señalaron la necesidad de incorporar "
        "herramientas digitales para un curso de negocios del siglo XXI."
    )

    write_paragraph(paras[97],
        "Desde el conectivismo (Siemens, 2005), el aprendizaje en la era "
        "digital requiere que los estudiantes interactúen con nodos de "
        "conocimiento distribuido y desarrollen la capacidad de identificar "
        "patrones en redes de información. Los recursos actuales del "
        "módulo operan en un paradigma predominantemente analógico que "
        "no aprovecha este potencial. Incorporar recursos digitales no es "
        "solo una actualización tecnológica, sino una necesidad pedagógica "
        "para formar profesionistas competentes en un entorno empresarial "
        "cada vez más mediado por datos y plataformas digitales."
    )


def fill_parte10_conclusiones(doc):
    """Parte 10 (P102): Conclusiones y recomendaciones."""
    paras = doc.paragraphs

    write_paragraph(paras[102],
        "A partir del análisis realizado, se concluye que la planeación del "
        "Módulo 1 del bloque MT1001B presenta fortalezas significativas: una "
        "estructura clara y bien secuenciada, un modelo pedagógico 5E adaptado "
        "que integra fases de empatía e innovación, competencias bien definidas "
        "con niveles de dominio diferenciados, y un balance adecuado entre "
        "teoría y práctica (calificación de 8.0/10 en el Delphi en este rubro). "
        "El uso de aprendizaje basado en retos y la vinculación con socios "
        "formadores son elementos valiosos del Modelo Tec21 que enriquecen la "
        "experiencia formativa.\n\n"

        "No obstante, el análisis revela cuatro áreas de oportunidad principales:\n\n"

        "1. Integración tecnológica insuficiente (5.5/10 Delphi): Las estrategias "
        "dependen de métodos presenciales tradicionales sin aprovechar "
        "herramientas digitales de análisis, visualización o colaboración.\n\n"

        "2. Evaluación poco diversificada: Predomina la evaluación sumativa "
        "de productos; falta evaluación diagnóstica explícita, instrumentos "
        "para competencias actitudinales y mecanismos de evaluación entre pares.\n\n"

        "3. Perspectiva crítica ausente: Siguiendo a Apple (1979), los contenidos "
        "no invitan a cuestionar las implicaciones éticas o sociales de las "
        "estrategias de mercadotecnia, ofreciendo una visión exclusivamente "
        "funcionalista.\n\n"

        "4. Autonomía estudiantil limitada: A pesar de que el tránsito de nivel "
        "A a B implica mayor agencia, la mayoría de actividades son guiadas por "
        "el docente, con pocas oportunidades para la investigación autónoma o "
        "la toma de decisiones independiente.\n\n"

        "Estas áreas de oportunidad fundamentan la propuesta de mejora que se "
        "presenta en la siguiente sección."
    )


def fill_parte11_propuesta(doc):
    """Parte 11 (P105-P107): Propuesta de mejora — Integración tecnológica."""
    paras = doc.paragraphs

    write_paragraph(paras[105],
        "Propuesta de Mejora: Plan de Integración Tecnológica para MT1001B",
        bold=True
    )

    write_paragraph(paras[106],
        "Con base en el hallazgo principal del análisis —la calificación de "
        "5.5/10 en integración tecnológica según la simulación Wideband "
        "Delphi—, se propone un plan de acción en tres fases para transformar "
        "la experiencia de aprendizaje del bloque MT1001B incorporando "
        "herramientas digitales de manera pedagógicamente significativa.\n\n"

        "Fase 1 — Diagnóstico digital (Semana previa al inicio del módulo)\n\n"

        "• Mapeo de herramientas: Realizar un inventario de las herramientas "
        "digitales actualmente disponibles en la institución (licencias, "
        "plataformas LMS, software de análisis) versus las herramientas "
        "deseadas para cada competencia del bloque.\n"
        "• Encuesta de competencias digitales: Aplicar un diagnóstico breve "
        "a los estudiantes para identificar su nivel de alfabetización "
        "digital y familiaridad con herramientas de análisis de datos.\n"
        "• Objetivo: Establecer una línea base que permita diseñar la "
        "integración tecnológica de forma diferenciada y realista.\n\n"

        "Fase 2 — Integración progresiva (Durante las 5 semanas del módulo)\n\n"

        "• Sesiones 1-4 (Diagnóstico): Incorporar Google Trends y SimilarWeb "
        "para el análisis de mercado del caso de estudio, permitiendo que "
        "el diagnóstico se base en datos reales, no solo en información "
        "del caso impreso.\n"
        "• Sesiones 5-7 (Conceptualización): Utilizar Canva o Piktochart "
        "para la elaboración de reportes visuales, fomentando las "
        "multiliteracidades (Cope & Kalantzis, 2009) mediante la "
        "representación multimodal del conocimiento.\n"
        "• Sesiones 8-10 (Estrategia e innovación): Integrar Miro o "
        "FigJam para el taller de design thinking colaborativo digital, "
        "y herramientas de visualización de datos (Tableau Public, "
        "Power BI) para la presentación del plan de mercadotecnia.\n"
        "• Transversal: Utilizar una plataforma colaborativa (Notion, "
        "Padlet) como espacio de trabajo compartido del equipo, "
        "alineándose con el conectivismo (Siemens, 2005) al crear "
        "redes de conocimiento entre los participantes.\n\n"

        "Fase 3 — Evaluación y sostenibilidad (Cierre del módulo y seguimiento)\n\n"

        "• Rúbricas digitales: Diseñar rúbricas en herramientas como "
        "CoRubrics que incluyan criterios de competencia digital y "
        "permitan la autoevaluación y evaluación entre pares.\n"
        "• Portafolio electrónico: Implementar un portafolio digital "
        "(Google Sites, Notion) donde los estudiantes documenten su "
        "proceso de aprendizaje, sus productos y sus reflexiones, "
        "fortaleciendo la metacognición.\n"
        "• Métricas de impacto: Definir indicadores para evaluar la "
        "efectividad de la integración tecnológica: satisfacción "
        "estudiantil, calidad de los productos, nivel de competencia "
        "digital alcanzado, y comparación con cohortes anteriores.\n"
        "• Capacitación docente: Asegurar que los profesores del bloque "
        "reciban formación en las herramientas seleccionadas, para que "
        "la integración no dependa de la iniciativa individual sino que "
        "se institucionalice.\n\n"

        "Esta propuesta busca elevar la calificación de integración "
        "tecnológica de 5.5/10 a al menos 7.5/10 en una implementación "
        "piloto, manteniendo las fortalezas ya identificadas (estructura "
        "clara, modelo 5E, balance teoría-práctica) y atendiendo la "
        "principal área de oportunidad de forma gradual y sostenible."
    )


def fill_parte12_adjunta(doc):
    """Parte 12 (P110-P111): Referencia a planeación adjunta."""
    paras = doc.paragraphs

    write_paragraph(paras[110],
        "La planeación analizada corresponde al documento institucional del "
        "bloque MT1001B «Descubrimientos del mercado para el desarrollo de "
        "estrategias» del Tecnológico de Monterrey. El documento original es "
        "un PDF escaneado de 37 páginas que fue transcrito en 6 archivos de "
        "referencia para este análisis:"
    )

    write_paragraph(paras[111],
        "• competencias.md — Competencias SNEG0400, SNEG0700, SEG0200 con "
        "subcompetencias y niveles A/B\n"
        "• objetivos_aprendizaje.md — 5 objetos de aprendizaje del Módulo 1 "
        "con contenidos procedimentales y actitudinales\n"
        "• organizacion_curso.md — Tabla detallada de 10 sesiones, 20h clase "
        "+ 10h tarea, 5 semanas\n"
        "• perfil_ingreso.md — Perfil de entrada: estudiantes de 2° semestre, "
        "Entrada de Negocios, niveles A-B\n"
        "• secuencia_didactica.md — Modelo 5E adaptado: Empatiza → Explica → "
        "Explora → Elabora → Innova → Evalúa\n"
        "• referencias_teoricas.md — Marcos teóricos: Multiliteracidades, "
        "Apple, Conectivismo, Sacristán"
    )


# ── Main ───────────────────────────────────────────────────────────

def main():
    template_path = "Análisis de Planeación.docx"
    output_path = "Análisis de Planeación - MT1001B.docx"

    print(f"Abriendo template: {template_path}")
    doc = Document(template_path)

    print("Llenando metadatos...")
    fill_metadata(doc)

    print("Llenando Parte 1: Contexto...")
    fill_parte1_contexto(doc)

    print("Llenando Parte 2: Datos de presentación...")
    fill_parte2_datos(doc)

    print("Llenando Parte 3: Propósito...")
    fill_parte3_proposito(doc)

    print("Llenando Parte 4: Competencias y aprendizajes...")
    fill_parte4_competencias(doc)

    print("Llenando Parte 5: Contenidos...")
    fill_parte5_contenidos(doc)

    print("Llenando Parte 6: Secuencia didáctica...")
    fill_parte6_secuencia(doc)

    print("Llenando Parte 7: Metodología...")
    fill_parte7_metodologia(doc)

    print("Llenando Parte 8: Evaluación...")
    fill_parte8_evaluacion(doc)

    print("Llenando Parte 9: Recursos...")
    fill_parte9_recursos(doc)

    print("Llenando Parte 10: Conclusiones...")
    fill_parte10_conclusiones(doc)

    print("Llenando Parte 11: Propuesta de mejora...")
    fill_parte11_propuesta(doc)

    print("Llenando Parte 12: Planeación adjunta...")
    fill_parte12_adjunta(doc)

    print(f"Guardando documento: {output_path}")
    doc.save(output_path)
    print("¡Listo! Documento generado exitosamente.")


if __name__ == "__main__":
    main()
